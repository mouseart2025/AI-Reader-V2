"""Series Bible Word (.docx) renderer — convert collected data to Word document.

Uses python-docx to generate a styled .docx with TOC placeholder, headers,
page numbers, and content mirroring the Markdown renderer.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from src.services.series_bible_service import SeriesBibleData


def render_docx(data: SeriesBibleData, template: str = "complete") -> io.BytesIO:
    """Render SeriesBibleData as a Word document. Returns BytesIO buffer."""
    doc = Document()

    # ── Page setup ────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Default font ──────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── Header / Footer ──────────────────────────
    for section in doc.sections:
        # Header: novel title
        header_para = section.header.paragraphs[0]
        header_label = "网文作者套件" if template == "author" else "设定集"
        header_para.text = f"{data.novel_title} — {header_label}"
        header_para.style = doc.styles["Header"]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.runs[0] if header_para.runs else header_para.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Footer: page number
        footer_para = section.footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(footer_para)

    # ── Title page ────────────────────────────────
    title_para = doc.add_heading(data.novel_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if data.novel_author:
        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_p.add_run(f"作者: {data.novel_author}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    range_p = doc.add_paragraph()
    range_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = range_p.add_run(f"分析范围: 第 {data.chapter_range[0]} ~ {data.chapter_range[1]} 章")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── TOC placeholder ──────────────────────────
    doc.add_paragraph()
    toc_heading = doc.add_heading("目录", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_toc_field(doc)
    doc.add_page_break()

    # ── Determine active modules ──────────────────
    modules = data.modules or [
        "characters", "relations", "locations", "items", "orgs", "timeline",
    ]

    # ── Characters ────────────────────────────────
    if "characters" in modules and data.characters:
        doc.add_heading("人物档案", level=1)
        _render_characters(doc, data.characters)

    # ── Relations ─────────────────────────────────
    if "relations" in modules and data.relations:
        doc.add_heading("关系网络", level=1)
        _render_relations(doc, data.relations)

    # ── Locations ─────────────────────────────────
    if "locations" in modules and data.locations:
        doc.add_heading("地点百科", level=1)
        _render_locations(doc, data.locations)

    # ── Items ─────────────────────────────────────
    if "items" in modules and data.items:
        doc.add_heading("物品道具", level=1)
        _render_items(doc, data.items)

    # ── Orgs ──────────────────────────────────────
    if "orgs" in modules and data.orgs:
        doc.add_heading("组织势力", level=1)
        _render_orgs(doc, data.orgs)

    # ── Timeline ──────────────────────────────────
    if "timeline" in modules and data.timeline:
        doc.add_heading("时间线", level=1)
        _render_timeline(doc, data.timeline)

    # ── Footer ────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("由 AI Reader V2 自动生成")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # ── Write to buffer ───────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Section renderers ─────────────────────────────


def _render_characters(doc: Document, characters: list[dict]) -> None:
    for ch in characters:
        doc.add_heading(ch["name"], level=2)

        aliases = ch.get("aliases", [])
        alias_names = [a["name"] for a in aliases if a["name"] != ch["name"]]
        if alias_names:
            _add_field(doc, "别称", ", ".join(alias_names))

        appearances = ch.get("appearances", [])
        if appearances:
            _add_field(doc, "外貌特征", "")
            for ap in appearances[:3]:
                doc.add_paragraph(ap["description"], style="List Bullet")

        abilities = ch.get("abilities", [])
        if abilities:
            _add_field(doc, "能力", "")
            for ab in abilities[:5]:
                dim = ab.get("dimension", "")
                name = ab.get("name", "")
                desc = ab.get("description", "")
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{dim}·{name}")
                run.bold = True
                p.add_run(f": {desc}")

        relations = ch.get("relations", [])
        if relations:
            _add_field(doc, "人物关系", "")
            for rel in relations[:10]:
                other = rel.get("other_person", "")
                category = rel.get("category", "other")
                stages = rel.get("stages", [])
                if len(stages) > 1:
                    chain = " → ".join(_compress_chain(stages))
                    doc.add_paragraph(
                        f"{other} — {chain} ({_cat_label(category)})",
                        style="List Bullet",
                    )
                elif stages:
                    rel_type = stages[0].get("relation_type", "")
                    doc.add_paragraph(
                        f"{other} — {rel_type} ({_cat_label(category)})",
                        style="List Bullet",
                    )
                else:
                    doc.add_paragraph(other, style="List Bullet")

        experiences = ch.get("experiences", [])
        if experiences:
            _add_field(doc, "主要经历", "")
            for exp in experiences[:8]:
                ch_num = exp.get("chapter", "")
                summary = exp.get("summary", "")
                loc = exp.get("location")
                loc_str = f" @ {loc}" if loc else ""
                doc.add_paragraph(
                    f"[第{ch_num}章] {summary}{loc_str}", style="List Bullet"
                )

        items = ch.get("items", [])
        if items:
            _add_field(doc, "持有物品", "")
            for it in items[:5]:
                doc.add_paragraph(
                    f"{it.get('item_name', '')} — {it.get('action', '')} (第{it.get('chapter', '')}章)",
                    style="List Bullet",
                )

        stats = ch.get("stats", {})
        if stats:
            p = doc.add_paragraph()
            run = p.add_run(
                f"出场章数: {stats.get('chapter_count', 0)} · "
                f"首次出场: 第{stats.get('first_chapter', '?')}章"
            )
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True


def _render_relations(doc: Document, relations: dict) -> None:
    edges = relations.get("edges", [])
    if not edges:
        doc.add_paragraph("暂无关系数据")
        return

    sorted_edges = sorted(edges, key=lambda e: e.get("weight", 0), reverse=True)
    display = sorted_edges[:30]

    table = doc.add_table(rows=1, cols=4, style="Table Grid")
    hdr = table.rows[0].cells
    for i, label in enumerate(["人物A", "人物B", "关系", "章节数"]):
        hdr[i].text = label
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for edge in display:
        row = table.add_row().cells
        row[0].text = edge.get("source", "")
        row[1].text = edge.get("target", "")
        row[2].text = edge.get("relation_type", "")
        row[3].text = str(edge.get("weight", 0))


def _render_locations(doc: Document, locations: list[dict]) -> None:
    for loc in locations:
        doc.add_heading(loc["name"], level=2)

        meta = []
        loc_type = loc.get("location_type", "")
        parent = loc.get("parent")
        if loc_type:
            meta.append(f"类型: {loc_type}")
        if parent:
            meta.append(f"上级: {parent}")
        if meta:
            _add_field(doc, " · ".join(meta), "")

        children = loc.get("children", [])
        if children:
            _add_field(doc, "下级地点", ", ".join(children[:10]))

        descriptions = loc.get("descriptions", [])
        if descriptions:
            for desc in descriptions[:3]:
                doc.add_paragraph(desc.get("description", ""), style="List Bullet")

        visitors = loc.get("visitors", [])
        if visitors:
            visitor_names = [
                f"{v['name']}{'(常驻)' if v.get('is_resident') else ''}"
                for v in visitors[:8]
            ]
            _add_field(doc, "到访者", ", ".join(visitor_names))

        stats = loc.get("stats", {})
        if stats:
            p = doc.add_paragraph()
            run = p.add_run(
                f"提及章数: {stats.get('chapter_count', 0)} · "
                f"首次出现: 第{stats.get('first_chapter', '?')}章"
            )
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.italic = True


def _render_items(doc: Document, items: list[dict]) -> None:
    for item in items:
        doc.add_heading(item["name"], level=2)

        item_type = item.get("item_type", "")
        if item_type:
            _add_field(doc, "类型", item_type)

        flow = item.get("flow", [])
        if flow:
            _add_field(doc, "流转记录", "")
            for f in flow[:8]:
                ch_num = f.get("chapter", "")
                action = f.get("action", "")
                actor = f.get("actor", "")
                desc = f.get("description", "")
                text = f"[第{ch_num}章] {actor} {action}"
                if desc:
                    text += f" — {desc}"
                doc.add_paragraph(text, style="List Bullet")


def _render_orgs(doc: Document, orgs: list[dict]) -> None:
    for org in orgs:
        doc.add_heading(org["name"], level=2)

        org_type = org.get("org_type", "")
        if org_type:
            _add_field(doc, "类型", org_type)

        members = org.get("member_events", [])
        if members:
            _add_field(doc, "成员变动", "")
            for m in members[:10]:
                ch_num = m.get("chapter", "")
                member = m.get("member", "")
                action = m.get("action", "")
                role = m.get("role", "")
                role_str = f" ({role})" if role else ""
                doc.add_paragraph(
                    f"[第{ch_num}章] {member}{role_str} — {action}",
                    style="List Bullet",
                )

        org_rels = org.get("org_relations", [])
        if org_rels:
            _add_field(doc, "组织关系", "")
            for r in org_rels[:5]:
                doc.add_paragraph(
                    f"{r.get('other_org', '')} — {r.get('relation_type', '')}",
                    style="List Bullet",
                )


def _render_timeline(doc: Document, events: list[dict]) -> None:
    current_chapter = -1
    for ev in events[:500]:
        ch_num = ev.get("chapter", 0)
        if ch_num != current_chapter:
            current_chapter = ch_num
            doc.add_heading(f"第{ch_num}章", level=3)

        importance = ev.get("importance", "medium")
        summary = ev.get("summary", "")
        participants = ev.get("participants", [])
        loc = ev.get("location")

        parts = []
        if participants:
            parts.append(f"[{', '.join(participants[:3])}]")
        if loc:
            parts.append(f"@ {loc}")
        context = f" ({' '.join(parts)})" if parts else ""

        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{summary}{context}")
        if importance == "high":
            run.bold = True


# ── Helpers ───────────────────────────────────────


def _compress_chain(stages: list[dict]) -> list[str]:
    """Remove consecutive duplicate relation types from a stage chain."""
    if not stages:
        return []
    result = [stages[0].get("relation_type", "")]
    for s in stages[1:]:
        rt = s.get("relation_type", "")
        if rt != result[-1]:
            result.append(rt)
    return result


def _add_field(doc: Document, label: str, value: str) -> None:
    """Add a bold-label + value paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label}")
    run.bold = True
    if value:
        p.add_run(f" {value}")


def _add_page_number(paragraph) -> None:
    """Insert a PAGE field into the paragraph for auto page numbers."""
    run = paragraph.add_run()
    fld_char_begin = _make_element("w:fldChar", {"w:fldCharType": "begin"})
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = _make_element("w:instrText")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = _make_element("w:fldChar", {"w:fldCharType": "end"})
    run3._r.append(fld_char_end)


def _add_toc_field(doc: Document) -> None:
    """Insert a TOC field that Word will update on open."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = _make_element("w:fldChar", {"w:fldCharType": "begin"})
    run._r.append(fld_begin)

    run2 = p.add_run()
    instr = _make_element("w:instrText", {"xml:space": "preserve"})
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instr)

    run3 = p.add_run()
    fld_separate = _make_element("w:fldChar", {"w:fldCharType": "separate"})
    run3._r.append(fld_separate)

    run4 = p.add_run("(请在 Word 中右键此处 → 更新域 以生成目录)")
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.size = Pt(9)

    run5 = p.add_run()
    fld_end = _make_element("w:fldChar", {"w:fldCharType": "end"})
    run5._r.append(fld_end)


def _make_element(tag: str, attrs: dict | None = None):
    """Create an OPC XML element with optional attributes."""
    from lxml import etree

    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    element = etree.SubElement(etree.Element("dummy", nsmap=nsmap), qn(tag))
    if attrs:
        for k, v in attrs.items():
            element.set(qn(k) if ":" in k else k, v)
    return element


def _cat_label(category: str) -> str:
    """Translate relation category to Chinese label."""
    return {
        "family": "亲属",
        "intimate": "亲密",
        "social": "社交",
        "hostile": "敌对",
        "hierarchical": "上下级",
        "other": "其他",
    }.get(category, category)
